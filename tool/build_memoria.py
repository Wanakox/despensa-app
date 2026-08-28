from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path('/tmp/despensa-template/Plantilla-Trabajo_Practico.docx')
OUT = ROOT / 'docs/entrega/Despensa-Memoria-Final.docx'
BLUE, PALE, GRAY = '17365D', 'F2F6FA', '666666'

def clear_body(doc):
    body, sect = doc._element.body, doc._element.body.sectPr
    for child in list(body):
        if child is not sect: body.remove(child)

def set_cell(cell, value, bold=False, color='000000', size=8.5):
    cell.text = ''; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(value)); r.bold = bold; r.font.name = 'Garamond'; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    mar = OxmlElement('w:tcMar')
    for side, width in [('top',70),('start',90),('bottom',70),('end',90)]:
        n=OxmlElement(f'w:{side}'); n.set(qn('w:w'),str(width)); n.set(qn('w:type'),'dxa'); mar.append(n)
    cell._tc.get_or_add_tcPr().append(mar)

def shade(cell, fill):
    n=OxmlElement('w:shd'); n.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(n)

def add_table(doc, headers, rows, widths=None, size=8.5):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    borders=OxmlElement('w:tblBorders')
    for edge in ['top','left','bottom','right','insideH','insideV']:
        n=OxmlElement(f'w:{edge}'); n.set(qn('w:val'),'single'); n.set(qn('w:sz'),'4'); n.set(qn('w:color'),'9FBAD0'); borders.append(n)
    t._tbl.tblPr.append(borders)
    flag=OxmlElement('w:tblHeader'); flag.set(qn('w:val'),'true'); t.rows[0]._tr.get_or_add_trPr().append(flag)
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,True,'FFFFFF',size); shade(t.rows[0].cells[i],BLUE)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell(cells[i],v,False,'000000',size)
            if ri%2: shade(cells[i],PALE)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def para(doc,text):
    p=doc.add_paragraph(style='Normal'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.add_run(text); return p
def bullet(doc,text): doc.add_paragraph(text,style='List Bullet')
def page(doc): doc.add_page_break()
def figure(doc,relative,caption,width=6.0):
    path=ROOT/relative
    if not path.exists(): return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; shape=p.add_run().add_picture(str(path),width=Inches(width))
    shape._inline.docPr.set('descr',caption); shape._inline.docPr.set('title',caption)
    c=doc.add_paragraph(caption,style='Caption'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER

doc=Document(TEMPLATE)
for section in doc.sections:
    for part in [section.footer, section.first_page_footer]:
        for paragraph in part.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace('Vuestro pie de página', 'Despensa · Memoria del proyecto')
clear_body(doc)
for name in ['Normal','Heading 1','Heading 2','Heading 3','Caption','List Bullet','List Bullet 2']:
    if name not in doc.styles: doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
if 'Table Grid' not in doc.styles:
    doc.styles.add_style('Table Grid', WD_STYLE_TYPE.TABLE)
doc.styles['Normal'].font.name='Garamond'; doc.styles['Normal'].font.size=Pt(11); doc.styles['Normal'].paragraph_format.space_after=Pt(6); doc.styles['Normal'].paragraph_format.line_spacing=1.15
for name,size,before,after in [('Heading 1',12,0,3),('Heading 2',11.5,3,2),('Heading 3',11,2,1)]:
    s=doc.styles[name]; s.font.name='Garamond'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(BLUE); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.left_indent=Inches(0); s.paragraph_format.first_line_indent=Inches(0); s.paragraph_format.keep_with_next=True
doc.styles['Caption'].font.name='Garamond'; doc.styles['Caption'].font.size=Pt(9); doc.styles['Caption'].font.italic=True; doc.styles['Caption'].font.color.rgb=RGBColor.from_string(GRAY)

# Portada institucional de la plantilla.
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(100)
p=doc.add_paragraph('MEMORIA DEL PROYECTO / TRABAJO',style='Heading 1'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18)
p=doc.add_paragraph('DESPENSA'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.runs[0]; r.font.name='Garamond'; r.font.size=Pt(23); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph('Aplicación móvil para el inventario doméstico y la cesta compartida'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.name='Garamond'; p.runs[0].font.size=Pt(14); p.runs[0].font.bold=True
figure(doc,'docs/branding/despensa-logo.png','Identidad visual de Despensa',2.2)
p=doc.add_paragraph('Autor: Juan García'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.bold=True
p=doc.add_paragraph('Córdoba, agosto de 2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
page(doc)

doc.add_heading('ÍNDICE DE CONTENIDOS',1)
toc = [
    (0,'1. Especificación y Análisis del problema'),
    (1,'1.1. Visión general del proyecto de aplicación móvil'),
    (2,'1.1.1. Presentación del proyecto'), (2,'1.1.2. Objetivos'),
    (2,'1.1.3. Público objetivo'), (2,'1.1.4. Objetivos cuantitativos'),
    (2,'1.1.5. Especificaciones técnicas'), (2,'1.1.6. Recursos existentes'),
    (2,'1.1.7. Aplicaciones similares'),
    (1,'1.2. Descripción gráfica y ergonómica de la aplicación'),
    (2,'1.2.1. Identidad gráfica'), (2,'1.2.2. Diseño de interfaz y navegación'),
    (1,'1.3. Descripción funcional y técnica. Métricas de rendimiento (KPI)'),
    (2,'1.3.1. Requisitos funcionales'), (2,'1.3.2. Requisitos no funcionales'),
    (2,'1.3.3. Modelo de casos de uso'), (2,'1.3.4. KPI'),
    (1,'1.4. Desarrollo backend de aplicaciones móviles'),
    (2,'1.4.1. Arquitectura'), (2,'1.4.2. Modelo de datos'),
    (2,'1.4.3. Seguridad y tiempo real'),
    (1,'1.5. Descripción de los tipos de pruebas a realizar'),
    (2,'1.5.1. Herramientas'), (2,'1.5.2. Alcance y metodología'),
    (2,'1.5.3. Resultado'),
    (1,'1.6. Servicios integrados, proveedores y planificación'),
    (0,'2. Desarrollo del proyecto acorde a SCRUM'),
    (1,'2.1. Herramientas para la gestión'), (1,'2.2. Roles'),
    (1,'2.3. Componentes y reuniones'), (1,'2.4. Sprints'),
    (2,'2.4.1. Sprint 1 — Análisis y Diseño'),
    (2,'2.4.2. Sprint 2 — Base técnica e inventario'),
    (2,'2.4.3. Sprint 3 — Cesta, caducidad y colaboración'),
    (2,'2.4.4. Sprint 4 — Calidad y cierre'),
    (1,'2.5. Seguimiento final de issues'),
    (0,'3. Casos de Prueba'), (1,'3.1. Especificación'),
    (0,'4. Conclusiones'), (1,'4.1. Puntos fuertes'),
    (1,'4.2. Puntos débiles'), (1,'4.3. Mejoras futuras'),
    (0,'Referencias bibliográficas'), (0,'ANEXOS'),
    (1,'ANEXO I. Especificaciones funcionales'),
    (1,'ANEXO II. Especificaciones formales del sistema'),
    (2,'1. Diseño arquitectónico'), (2,'2. Diseño de la base de datos'),
    (2,'3. Evidencias de interfaz y proceso'),
]
for level,item in toc:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.22*level); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0
    r=p.add_run(item); r.font.name='Garamond'; r.font.size=Pt(9.5 if level==0 else 9); r.bold=(level==0); r.font.color.rgb=RGBColor.from_string(BLUE)
page(doc)

doc.add_heading('1. Especificación y Análisis del problema',1)
para(doc,'Despensa nace para resolver una dificultad cotidiana: en los hogares compartidos, la información sobre lo que queda, lo que va a caducar y lo que debe comprarse suele repartirse entre la memoria de cada persona, notas y conversaciones. La falta de una fuente común provoca compras duplicadas, productos agotados sin previsión y alimentos desperdiciados.')
doc.add_heading('1.1. Visión general del proyecto de aplicación móvil',2)
doc.add_heading('1.1.1. Presentación del proyecto',3)
para(doc,'Despensa es una aplicación móvil colaborativa para Android que centraliza el inventario de un hogar y lo conecta con una cesta común. Cada usuario accede a uno o varios hogares, consulta existencias, registra cantidades y caducidades y coordina compras con el resto de miembros. La propuesta de valor es hacer explícito y breve el ciclo entre existencias, reposición, compra e incorporación al inventario.')
doc.add_heading('1.1.2. Objetivos',3)
for t in ['Centralizar productos, cantidades, unidades, categorías, ubicaciones y caducidades.','Mantener una cesta sincronizada entre miembros.','Evitar duplicados al enviar productos a la cesta o incorporar una compra.','Destacar agotados, caducados o próximos a caducar con margen configurable.','Proteger los datos mediante autenticación, pertenencia y roles.','Ofrecer una interfaz clara y un modo local de demostración.']: bullet(doc,t)
doc.add_heading('1.1.3. Público objetivo',3)
para(doc,'La aplicación se dirige a familias, parejas, pisos de estudiantes y grupos que compartan productos y compras. El usuario objetivo utiliza Android, necesita introducir información con rapidez y valora observar los cambios de otras personas sin coordinarse por canales externos.')
doc.add_heading('1.1.4. Objetivos cuantitativos',3)
add_table(doc,['Indicador','Objetivo','Resultado'],[('Cobertura','Completar los 16 casos principales','Implementados'),('Calidad','Cero incidencias estáticas','Cumplido'),('Pruebas','Superar la batería automatizada','25/25'),('Trazabilidad','Milestones e issues','4 milestones y 45 issues'),('Entrega','APK Android','Release de 55,2 MB')],[1.4,2.4,2.5])
doc.add_heading('1.1.5. Especificaciones técnicas',3)
for t in ['Flutter y Dart.','Android como plataforma inicial.','Firebase Authentication y Cloud Firestore.','SharedPreferences para modo local.','Material 3 y diseño adaptable.','Arquitectura organizada por funcionalidades.']: bullet(doc,t)
doc.add_heading('1.1.6. Recursos existentes',3)
para(doc,'El proyecto utiliza Visual Studio Code, Flutter SDK, Android SDK, Firebase Console, Git y GitHub. Dispone de identidad gráfica, diagramas, documentos de visión, requisitos, casos de uso y datos, capturas Scrum y pruebas Flutter.')
doc.add_heading('1.1.7. Aplicaciones similares',3)
para(doc,'El análisis de docs/analisis-aplicaciones-similares.odt estudia NoWaste, Pantry Check, Bring! y Listonic. Las primeras destacan por inventario y caducidades; las últimas, por listas compartidas. Despensa une ambos extremos sin añadir recetas, ofertas, precios o inteligencia artificial al MVP.')
add_table(doc,['Aplicación','Fortaleza','Decisión para Despensa'],[('NoWaste','Caducidades','Entrada sencilla'),('Pantry Check','Inventario y escaneo','Escaneo como mejora'),('Bring!','Lista compartida','Cesta con pocos pasos'),('Listonic','Sincronización','Permisos mediante cuentas')],[1.2,2.3,2.8])

doc.add_heading('1.2. Descripción gráfica y ergonómica de la aplicación',2)
doc.add_heading('1.2.1. Identidad gráfica',3)
para(doc,'La identidad utiliza verde salvia como color principal, marfil como fondo y terracota para avisos. El símbolo combina una casa y un recipiente, relacionando hogar y almacenamiento. Las formas redondeadas transmiten cercanía manteniendo legibilidad.')
figure(doc,'docs/branding/despensa-logo.png','Figura 1. Logotipo principal de Despensa',2.6)
doc.add_heading('1.2.2. Diseño de interfaz y navegación',3)
para(doc,'La navegación separa la selección del hogar del trabajo dentro de este. La barra inferior mantiene accesibles Inicio, Despensa, Cesta y Miembros. Las operaciones destructivas solicitan confirmación; los formularios validan datos y los estados vacíos indican la acción siguiente. Iconos, etiquetas y tooltips reducen ambigüedad.')
figure(doc,'docs/capturas/sprint3-final.png','Figura 2. Estado funcional al cierre del Sprint 3',6.1)
figure(doc,'docs/capturas/sprint4-final.png','Figura 3. Tablero al cierre del proyecto',6.1)

doc.add_heading('1.3. Descripción funcional y técnica. Métricas de rendimiento (KPI)',2)
doc.add_heading('1.3.1. Requisitos funcionales',3)
reqs=[('RF-01','Registro, inicio y cierre de sesión.'),('RF-02','Crear y seleccionar hogares.'),('RF-03','Invitar, aceptar, rechazar y retirar miembros.'),('RF-04','Añadir, consultar, editar y eliminar productos.'),('RF-05','Buscar y filtrar el inventario.'),('RF-06','Mostrar agotamiento y caducidad.'),('RF-07','Configurar el margen de aviso.'),('RF-08','Gestionar elementos de cesta.'),('RF-09','Enviar productos a cesta sin duplicados.'),('RF-10','Incorporar compras al inventario.'),('RF-11','Sincronizar cambios en tiempo real.'),('RF-12','Mostrar actividad reciente.')]
add_table(doc,['ID','Requisito'],reqs,[0.8,5.5])
doc.add_heading('1.3.2. Requisitos no funcionales',3)
add_table(doc,['Área','Requisito'],[('Usabilidad','Flujos frecuentes breves y mensajes comprensibles.'),('Accesibilidad','Contraste y controles táctiles amplios.'),('Rendimiento','Actualizaciones sin bloquear la interfaz.'),('Seguridad','Acceso exclusivo de miembros autenticados.'),('Disponibilidad','Modo local para demostración.'),('Compatibilidad','Adaptación a distintas pantallas Android.'),('Mantenibilidad','Separación por funcionalidades y servicios.')],[1.3,5.0])
doc.add_heading('1.3.3. Modelo de casos de uso',3)
para(doc,'El catálogo de docs/casos-de-uso-despensa.odt define usuarios no registrados, usuarios autenticados, propietarios y Firebase Authentication. Cada caso incluye precondiciones, flujo principal, alternativas, errores y postcondiciones.')
figure(doc,'docs/diagrams/Despensa-CU.png','Figura 4. Diagrama general de casos de uso',6.1)
doc.add_heading('1.3.4. KPI',3)
add_table(doc,['KPI','Definición','Medición inicial'],[('Cobertura','Casos principales implementados','16'),('Estabilidad','Pruebas superadas','25/25'),('Calidad estática','Incidencias de analyze','0'),('Sincronización','Cambios entre miembros','Snapshots de Firestore'),('Desperdicio evitado','Productos atendidos antes de caducar','Medición futura')],[1.25,3.1,2.0])

doc.add_heading('1.4. Desarrollo backend de aplicaciones móviles',2)
doc.add_heading('1.4.1. Arquitectura',3)
para(doc,'Firebase actúa como Backend as a Service. Authentication mantiene la identidad y Cloud Firestore almacena hogares y subcolecciones. Los servicios Dart encapsulan consultas para desacoplar las pantallas. Cuando Firebase no está disponible, la misma interfaz usa SharedPreferences.')
doc.add_heading('1.4.2. Modelo de datos',3)
figure(doc,'docs/diagrams/Despensa-ER.png','Figura 5. Modelo entidad-relación',6.1)
add_table(doc,['Entidad','Responsabilidad'],[('Usuario','Identidad y referencias de pertenencia.'),('Hogar','Contexto, propietario, miembros y preferencias.'),('Producto','Existencias, ubicación y caducidad.'),('ElementoCesta','Producto pendiente o comprado.'),('Invitación','Destinatario, hogar y estado.'),('Actividad','Acción, autor y fecha de servidor.')],[1.55,4.75])
doc.add_heading('1.4.3. Seguridad y tiempo real',3)
para(doc,'Las reglas comprueban autenticación y pertenencia. Las acciones administrativas requieren rol owner. Aceptar una invitación valida el correo del token. Inventario, cesta, miembros y actividad se observan mediante streams de snapshots.')

doc.add_heading('1.5. Descripción de los tipos de pruebas a realizar',2)
doc.add_heading('1.5.1. Herramientas',3)
for t in ['flutter_test para pruebas unitarias y widgets.','flutter analyze para análisis estático.','Firebase Emulator Suite para reglas.','Gradle release para el APK.','Renderizado para revisar memoria y presentación.']: bullet(doc,t)
doc.add_heading('1.5.2. Alcance y metodología',3)
para(doc,'Las pruebas cubren persistencia, hogares, inventario, cantidades, búsqueda, filtros, cesta, compra, miembros y preferencias. La revisión de interfaz considera navegación, formularios, confirmaciones, mensajes y estados vacíos. Antes del APK se ejecutaron dependencias, análisis, pruebas y build release.')
doc.add_heading('1.5.3. Resultado',3)
para(doc,'El cierre obtuvo cero incidencias de análisis, 25 pruebas superadas y un APK de 55,2 MB. Su SHA-256 es bc62374731cfb7c337fcfcdc5d04f9dd855fb698cdd4d2a2d57177f8d74f2954.')

doc.add_heading('1.6. Servicios integrados, proveedores y planificación',2)
add_table(doc,['Servicio','Proveedor','Justificación'],[('Framework','Flutter / Google','Una base de código y Material 3.'),('Autenticación','Firebase Auth','SDK oficial.'),('Datos','Cloud Firestore','Tiempo real y reglas.'),('Local','SharedPreferences','Persistencia ligera.'),('Gestión','GitHub','Código, issues, milestones y release.')],[1.2,1.7,3.4])
add_table(doc,['Sprint','Objetivo','Entregable'],[('1','Análisis y diseño','Visión, requisitos, casos y datos.'),('2','Base e inventario','Flutter, Firebase, hogares e inventario.'),('3','Cesta y colaboración','Compra, avisos e invitaciones.'),('4','Calidad y cierre','Pruebas, APK, memoria y presentación.')],[0.7,2.4,3.2])

doc.add_heading('2. Desarrollo del proyecto acorde a SCRUM',1)
doc.add_heading('2.1. Herramientas para la gestión',2)
para(doc,'GitHub centraliza código y gestión. Las milestones representan los sprints; las issues se etiquetan por análisis, documentación, diseño, frontend, backend o base de datos. Commits y cierres conservan la trazabilidad hasta v1.0.0.')
for img,cap in [('docs/capturas/creacion_proyecto.png','Figura 6. Creación del proyecto'),('docs/capturas/labels_creados.png','Figura 7. Etiquetas de trabajo'),('docs/capturas/milestone_creados.png','Figura 8. Milestones'),('docs/capturas/issues_creadas.png','Figura 9. Product Backlog inicial')]: figure(doc,img,cap,6.1)
doc.add_heading('2.2. Roles',2)
para(doc,'Al ser un proyecto individual, Juan García asume Product Owner, Scrum Master y desarrollador: prioriza valor, mantiene cadencia e implementa, prueba y documenta el incremento.')
doc.add_heading('2.3. Componentes y reuniones',2)
for t in ['Product Backlog: issues del proyecto.','Sprint Backlog: issues de cada milestone.','Incremento: versión funcional al cierre.','Definition of Done: integración, análisis, pruebas y evidencia.']: bullet(doc,t)
para(doc,'Cada sprint tuvo planificación, seguimiento del tablero, revisión funcional y retrospectiva. Esta última ajustó alcance, trazabilidad y calidad para la siguiente iteración.')
doc.add_heading('2.4. Sprints',2)
sprints=[('Sprint 1 — Análisis y Diseño','Visión, alcance, competidores, requisitos, casos y datos.','docs/capturas/sprint1-estado_inicial.png','docs/capturas/sprint1-estado_final.png'),('Sprint 2 — Base técnica e inventario','Flutter, Firebase, autenticación, hogares, navegación e inventario.','docs/capturas/sprint2-issues.png','docs/capturas/sprint2-final.png'),('Sprint 3 — Cesta, caducidad y colaboración','Compra, avisos configurables, tiempo real e invitaciones.','docs/capturas/sprint3-inicio.png','docs/capturas/sprint3-final.png'),('Sprint 4 — Calidad y cierre','Pruebas, accesibilidad, APK, memoria, presentación y release.','docs/capturas/sprint4-inicio.png','docs/capturas/sprint4-final.png')]
for i,(title,body,start,end) in enumerate(sprints,1):
    doc.add_heading(f'2.4.{i}. {title}',3); para(doc,body); figure(doc,start,f'Figura {9+i*2}. Inicio del {title.split(" — ")[0]}',5.9); figure(doc,end,f'Figura {10+i*2}. Cierre del {title.split(" — ")[0]}',5.9)
doc.add_heading('2.5. Seguimiento final de issues',2)
para(doc,'La última revisión comprobó que las nueve tareas de calidad y cierre estaban documentadas y resueltas antes de cerrar la milestone y publicar la release.')
figure(doc,'docs/capturas/issues-4.png','Figura 18. Issues asociadas al Sprint 4',6.1)

doc.add_heading('3. Casos de Prueba',1)
doc.add_heading('3.1. Especificación',2)
cases=[('CP-01','Registro e inicio','Cuenta válida','Acceso a hogares','Superado'),('CP-02','Crear hogar','Nombre válido','Hogar creado','Superado'),('CP-03','Invitar miembro','Correo válido','Invitación','Superado'),('CP-04','Añadir producto','Formulario válido','Persistido','Superado'),('CP-05','Editar/eliminar','Producto existente','Cambio confirmado','Superado'),('CP-06','Buscar/filtrar','Consulta y filtro','Lista correcta','Superado'),('CP-07','Cantidad','Incremento/reducción','Actualizada','Superado'),('CP-08','Caducidad','Margen 7 días','Aviso aplicado','Superado'),('CP-09','Enviar a cesta','Agotado','Sin duplicados','Superado'),('CP-10','Editar cesta','Elemento existente','Actualizado','Superado'),('CP-11','Marcar comprado','Checkbox','Estado conservado','Superado'),('CP-12','Guardar compra','Comprados','Inventario actualizado','Superado'),('CP-13','Vaciar cesta','Confirmación','Estado vacío','Superado'),('CP-14','Permisos','Usuario ajeno','Acceso denegado','Revisado'),('CP-15','Análisis','flutter analyze','0 incidencias','Superado'),('CP-16','Release','Build Android','APK generado','Superado')]
add_table(doc,['ID','Caso','Datos','Resultado','Estado'],cases,[0.55,1.35,1.65,2.05,0.75],8)
para(doc,'La ejecución final de flutter test informó “All tests passed” con 25 pruebas. flutter analyze terminó sin incidencias. La presentación no presentó desbordamientos y la memoria fue renderizada página por página.')

doc.add_heading('4. Conclusiones',1)
doc.add_heading('4.1. Puntos fuertes',2)
for t in ['Ciclo completo inventario-cesta.','Colaboración en tiempo real y permisos.','Interfaz coherente.','Arquitectura por funcionalidades.','Trazabilidad hasta la release.']: bullet(doc,t)
doc.add_heading('4.2. Puntos débiles',2)
for t in ['Primera versión solo Android.','Sin notificaciones push.','Pruebas de reglas ampliables con Emulator Suite.','Registro manual costoso para inventarios grandes.']: bullet(doc,t)
doc.add_heading('4.3. Mejoras futuras',2)
for t in ['Códigos de barras.','Notificaciones locales y push.','Offline avanzado.','Estadísticas de desperdicio.','iOS y publicación.']: bullet(doc,t)
para(doc,'Despensa cumple el objetivo académico: una aplicación móvil completa, colaborativa, con backend real, pruebas, documentación y APK reproducible. La versión v1.0.0 es una base estable para evolucionar el producto.')

doc.add_heading('Referencias bibliográficas',1)
for ref in ['[1] Flutter. https://docs.flutter.dev/','[2] Firebase. https://firebase.google.com/docs','[3] Material Design 3. https://m3.material.io/','[4] Schwaber y Sutherland. The Scrum Guide, 2020.','[5] Universidad de Córdoba. Material de Ingeniería de Sistemas Móviles, 2026.','[6] NoWaste, Pantry Check, Bring! y Listonic. Información pública, agosto de 2026.']: para(doc,ref)

doc.add_heading('ANEXOS',1)
doc.add_heading('ANEXO I. Especificaciones funcionales',2)
use_cases=[('CU-01','Registrarse','No registrado'),('CU-02','Iniciar sesión','Usuario'),('CU-03','Cerrar sesión','Usuario'),('CU-04','Crear hogar','Usuario'),('CU-05','Consultar hogar','Miembro'),('CU-06','Invitar miembro','Propietario'),('CU-07','Retirar miembro','Propietario'),('CU-08','Añadir producto','Miembro'),('CU-09','Modificar producto','Miembro'),('CU-10','Eliminar producto','Miembro'),('CU-11','Buscar productos','Miembro'),('CU-12','Consultar cesta','Miembro'),('CU-13','Añadir a cesta','Miembro'),('CU-14','Modificar cesta','Miembro'),('CU-15','Retirar de cesta','Miembro'),('CU-16','Marcar comprado','Miembro')]
add_table(doc,['ID','Nombre','Actor'],use_cases,[0.9,3.2,2.15])
para(doc,'La especificación completa procede de docs/casos-de-uso-despensa.odt e incluye precondiciones, flujo principal, alternativas, errores y postcondiciones.')
figure(doc,'docs/diagrams/Despensa-CU.png','Anexo I — Figura A1. Casos de uso',6.1)
doc.add_heading('ANEXO II. Especificaciones formales del sistema',2)
doc.add_heading('1. Diseño arquitectónico',3)
for t in ['app/: aplicación y tema.','core/: servicios y widgets.','features/auth/: autenticación.','features/households y members/: hogares y permisos.','features/inventory y cart/: núcleo funcional.','features/home, activity y profile/: resumen, historial y cuenta.']: bullet(doc,t)
doc.add_heading('2. Diseño de la base de datos',3)
para(doc,'El diseño procede de docs/modelo-datos-despensa.odt. Hogar es el agregado principal; Miembro relaciona Usuario y Hogar; Producto y ElementoCesta pertenecen al hogar; Invitación controla la incorporación y Actividad registra cambios.')
figure(doc,'docs/diagrams/Despensa-ER.png','Anexo II — Figura A2. Modelo de datos',6.1)
doc.add_heading('3. Evidencias de interfaz y proceso',3)
for img,cap in [('docs/capturas/sprint1-estado_final.png','Anexo II — Figura A3. Sprint 1'),('docs/capturas/sprint2-final.png','Anexo II — Figura A4. Sprint 2'),('docs/capturas/sprint3-final.png','Anexo II — Figura A5. Sprint 3'),('docs/capturas/sprint4-final.png','Anexo II — Figura A6. Sprint 4')]: figure(doc,img,cap,6.1)

for p in doc.paragraphs:
    p.paragraph_format.widow_control=True
    for r in p.runs:
        if not r.font.name: r.font.name='Garamond'
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
